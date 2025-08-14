import streamlit as st
import openai
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile

# OpenRouter API config
openai.api_key = "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"
openai.api_base = "https://openrouter.ai/api/v1"

def split_lines(text):
    return [line.strip() for line in text.split('\n') if line.strip()]

# Streamlit UI
st.title("AI Resume Builder")

# Inputs
name = st.text_input("Full Name")
location = st.text_input("Location")
email = st.text_input("Email")
address = st.text_area("Residential Address")
linkedin = st.text_input("LinkedIn URL")
github = st.text_input("GitHub URL")
skills = st.text_area("Skills (comma separated)")
experience = st.text_area("Experience")
education = st.text_area("Education & Qualifications")
references = st.text_area("References")
hobbies = st.text_area("Hobbies")
manual_summary = st.text_area("Custom Summary (optional)")

auto_summary = ""

# AI Summary generation
if st.button("Generate Summary with AI"):
    if not name or not skills or not experience:
        st.error("Please enter at least Name, Skills, and Experience for AI summary.")
    else:
        prompt = f"Write a professional resume summary for {name}, located in {location}. Skills: {skills}. Experience: {experience}."
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        auto_summary = response.choices[0].message["content"]
        st.subheader("Generated Summary")
        st.write(auto_summary)

summary_to_use = manual_summary if manual_summary else auto_summary

# Show final resume preview and export
if st.button("Show Final Resume"):
    st.subheader("Resume Preview")
    st.write(f"**Name:** {name}")
    st.write(f"**Location:** {location}")
    st.write(f"**Email:** {email}")
    st.write(f"**Address:** {address}")
    st.write(f"**LinkedIn:** {linkedin}")
    st.write(f"**GitHub:** {github}")
    st.write(f"**Skills:** {skills}")
    st.write(f"**Experience:** {experience}")
    st.write(f"**Education:** {education}")
    st.write(f"**References:** {references}")
    st.write(f"**Hobbies:** {hobbies}")
    st.write(f"**Summary:** {summary_to_use}")

    # Prepare resume content for exports
    sections = [
        "Personal Information",
        "Online Profiles",
        "Skills",
        "Experience",
        "Education & Qualifications",
        "References",
        "Hobbies",
        "Summary"
    ]

    content = [
    f"Name: {name}\nLocation: {location}\nEmail: {email}\nAddress: {address}",
    f"LinkedIn: {linkedin}\nGitHub: {github}",
    "\n".join(f"- {s.strip()}" for s in skills.split(',')),
    "\n".join(f"- {line}" for line in split_lines(experience)),
    "\n".join(f"- {line}" for line in split_lines(education)),
    "\n".join(f"- {line}" for line in split_lines(references)),
    "\n".join(f"- {line}" for line in split_lines(hobbies)),
    summary_to_use
]

    # --- PDF export with styling ---
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 18)
            self.cell(0, 10, f'Resume - {name}', align='C', ln=True)
            self.ln(5)

    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    for header, text in zip(sections, content):
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 8, header, ln=True)
        pdf.ln(2)
        pdf.set_font("Arial", "", 12)
        for line in text.split('\n'):
            if line.startswith("•"):
                pdf.cell(10)  # indent bullet
                pdf.cell(0, 7, line, ln=True)
            else:
                pdf.multi_cell(0, 7, line)
        pdf.ln(4)

    pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(pdf_file.name)

    # --- DOCX export with styled headings & bullets ---
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    title = doc.add_heading(f"Resume - {name}", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for header, text in zip(sections, content):
        doc.add_heading(header, level=1)
        for line in text.split('\n'):
            if line.startswith("•"):
                para = doc.add_paragraph(line.lstrip("• ").strip(), style='List Bullet')
            else:
                para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)

    docx_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(docx_file.name)

    # --- HTML export with basic styling ---
    html_content = f"""
    <html>
    <head>
    <title>Resume - {name}</title>
    <style>
    body {{ font-family: Arial, sans-serif; margin: 30px; }}
    h1 {{ text-align: center; }}
    h2 {{ border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
    ul {{ list-style-type: disc; margin-left: 20px; }}
    </style>
    </head>
    <body>
    <h1>Resume - {name}</h1>
    """

    for header, text in zip(sections, content):
        html_content += f"<h2>{header}</h2>\n<ul>\n"
        for line in text.split('\n'):
            clean_line = line.lstrip("• ").strip()
            html_content += f"<li>{clean_line}</li>\n"
        html_content += "</ul>\n"

    html_content += "</body></html>"

    html_file = tempfile.NamedTemporaryFile(delete=False, suffix=".html")
    with open(html_file.name, "w", encoding="utf-8") as f:
        f.write(html_content)

    # Download buttons for all formats
    st.download_button("Download PDF", open(pdf_file.name, "rb").read(), file_name="resume.pdf", mime="application/pdf")
    st.download_button("Download DOCX", open(docx_file.name, "rb").read(), file_name="resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("Download HTML", open(html_file.name, "rb").read(), file_name="resume.html", mime="text/html")
